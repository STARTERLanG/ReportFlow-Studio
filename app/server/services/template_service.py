import json
from typing import Any

from docx import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_openai import ChatOpenAI

from agents.prompts.library import TEMPLATE_STRUCTURE_ANALYSIS_PROMPT
from app.server.config import settings
from app.server.logger import logger


class TemplateService:
    def __init__(self):
        logger.info(f"初始化 TemplateService LLM: Model={settings.llm.model_name}, BaseURL={settings.llm.base_url}")
        self.llm = ChatOpenAI(
            model=settings.llm.model_name,
            api_key=settings.llm.api_key,
            base_url=settings.llm.base_url,
            temperature=0,
            timeout=120,
        )

    def parse_and_decompose(self, file_path: str) -> list[dict[str, Any]]:
        """
        读取 Word 文档全文，并调用 AI 进行结构化分析（变量 + 任务）。
        """
        logger.info(f"正在读取模板文件: {file_path}")

        # 1. 结构化提取 (Markdown 风格)
        structured_content = self._extract_content_as_markdown(file_path)

        # 2. 调用 AI 进行分析
        logger.info(f"提取结构化文本成功 (长度: {len(structured_content)} 字符)，正在进行 AI 语义分析...")
        result = self._analyze_structure_with_llm(structured_content)

        return result

    def _extract_content_as_markdown(self, file_path: str) -> str:
        """
        将 Word 文档转换为 Markdown 格式，保留标题层级，以便 LLM 理解结构。
        """
        doc = Document(file_path)
        lines = []

        for element in doc.element.body:
            # 简单判断 tag 类型
            if element.tag.endswith("p"):  # 段落
                # 找到对应的 Paragraph 对象 (python-docx 没有直接从 element 反查 paragraph 的简单方法，这里简化处理)
                # 为保证顺序，重新遍历 paragraphs 和 tables 是比较麻烦的，
                # 这里我们简化：按顺序遍历 doc.iter_inner_content() (需要 python-docx >= 0.8.11)
                pass

        # 重新实现：按顺序遍历所有块级元素
        # python-docx 的 iter_block_items() 可以按顺序获取 paragraph 和 table
        # 这是一个辅助函数，用来按文档顺序生成 markdown

        def iter_block_items(parent):
            """
            Generate a reference to each paragraph and table child within *parent*,
            in document order. Each returned value is an instance of either Table or
            Paragraph.
            """
            if isinstance(parent, Document):
                parent_elm = parent.element.body
            else:
                parent_elm = parent._element

            for child in parent_elm.iterchildren():
                if child.tag.endswith("p"):
                    yield ("paragraph", child)
                elif child.tag.endswith("tbl"):
                    yield ("table", child)

        # 映射 Paragraph
        # 注意：这里需要一点 trick 把 element 映射回 Paragraph 对象，
        # 但为了简单和稳健，我们直接遍历 doc.paragraphs 和 doc.tables 可能会丢失相对顺序。
        # 更好的方法是：只根据 style.name 判断级别

        # 这种混合遍历比较复杂，我们采用一种更直接的策略：
        # 遍历 document.element.body 的子节点

        for block in self._iter_doc_blocks(doc):
            if block["type"] == "paragraph":
                p = block["obj"]
                text = p.text.strip()
                if not text:
                    continue

                style_name = p.style.name
                if style_name.startswith("Heading 1"):
                    lines.append(f"# {text}")
                elif style_name.startswith("Heading 2"):
                    lines.append(f"## {text}")
                elif style_name.startswith("Heading 3"):
                    lines.append(f"### {text}")
                else:
                    lines.append(text)

            elif block["type"] == "table":
                table = block["obj"]
                # 转换为简单的 Markdown 表格或 Key-Value 对
                lines.append("\n[表格数据开始]")
                for row in table.rows:
                    cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                    lines.append("| " + " | ".join(cells) + " |")
                lines.append("[表格数据结束]\n")

        return "\n\n".join(lines)

    def _iter_doc_blocks(self, doc):
        """
        按顺序遍历文档中的段落和表格
        """
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        parent = doc._body
        for child in parent._element.iterchildren():
            if isinstance(child, CT_P):
                yield {"type": "paragraph", "obj": Paragraph(child, parent)}
            elif isinstance(child, CT_Tbl):
                yield {"type": "table", "obj": Table(child, parent)}

    def _analyze_structure_with_llm(self, content: str) -> list[dict[str, Any]]:
        # 截断以防止 Token 溢出
        safe_content = content[:20000]

        prompt = ChatPromptTemplate.from_template(TEMPLATE_STRUCTURE_ANALYSIS_PROMPT)
        chain = prompt | self.llm

        try:
            response = chain.invoke({"content": safe_content})
            json_str = response.content.replace("```json", "").replace("```", "").strip()
            result = json.loads(json_str)

            # 兼容旧接口：如果 LLM 返回了 tasks 和 variables，直接返回 task 列表
            # 但我们需要把 variables 传出去，所以这里需要调整一下逻辑
            # 为了不破坏 parse_and_decompose 的签名 (list[dict])，我们可能需要 hack 一下
            # 或者，我们应该修改 parse_and_decompose 的返回类型？
            # 既然上一步 Schema 已经改了，这里直接返回 dict 最好，但在 API 层需要适配。
            # 为了最小改动，我们把 variables 附着在第一个 task 上？不，这太丑陋了。
            # 让我们直接返回 raw dict (包含 variables 和 tasks)，让 API 层去处理。

            # 直接返回任务列表，变量现在是任务内部的属性
            tasks = result.get("tasks", [])

            return {
                "variables": [],  # 废弃顶层 variables
                "tasks": tasks,
            }

        except Exception as e:
            logger.exception("AI 结构分析失败")
            return {"variables": [], "tasks": [{"task_name": "解析失败", "description": str(e), "requirements": ""}]}
