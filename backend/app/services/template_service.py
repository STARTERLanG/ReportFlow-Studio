import json
from typing import Any

from docx import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_openai import ChatOpenAI

from backend.agents.prompts.library import REPORT_TASK_DECOMPOSITION_PROMPT
from backend.app.config import settings
from backend.app.logger import logger


class TemplateService:
    def __init__(self):
        logger.info(f"初始化 TemplateService LLM: Model={settings.llm.model_name}, BaseURL={settings.llm.base_url}")
        self.llm = ChatOpenAI(
            model=settings.llm.model_name,
            api_key=settings.llm.api_key,
            base_url=settings.llm.base_url,
            temperature=0,
        )

    def parse_and_decompose(self, file_path: str) -> list[dict[str, Any]]:
        """
        读取 Word 文档全文，并调用 AI 进行任务拆解。
        """
        logger.info(f"正在读取模板文件: {file_path}")

        # 1. 简单暴力的全文提取
        full_text = self._extract_full_text(file_path)

        # 2. 调用 AI 进行拆解
        logger.info(f"提取文本成功 (长度: {len(full_text)} 字符)，正在进行 AI 任务拆解...")
        tasks = self._decompose_with_llm(full_text)

        return tasks

    def _extract_full_text(self, file_path: str) -> str:
        doc = Document(file_path)
        texts = []

        # 提取段落
        for para in doc.paragraphs:
            if para.text.strip():
                texts.append(para.text.strip())

        # 提取表格（简单按行拼接）
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                # 去重
                unique_cells = []
                seen = set()
                for t in row_cells:
                    if t not in seen:
                        unique_cells.append(t)
                        seen.add(t)
                if unique_cells:
                    texts.append(" | ".join(unique_cells))

        return "\n".join(texts)

    def _decompose_with_llm(self, content: str) -> list[dict[str, Any]]:
        # 截断以防止 Token 溢出 (保留前 15000 字符通常足够包含大纲)
        # 如果是超长文档，可能需要分片摘要，但 POC 阶段先这样
        safe_content = content[:15000]

        prompt = ChatPromptTemplate.from_template(REPORT_TASK_DECOMPOSITION_PROMPT)
        chain = prompt | self.llm

        try:
            response = chain.invoke({"content": safe_content})
            json_str = response.content.replace("```json", "").replace("```", "").strip()
            raw_tasks = json.loads(json_str)

            # 容错处理：修复 AI 可能出现的拼写错误
            cleaned_tasks = []
            for t in raw_tasks:
                # 修复 task_name 拼写错误
                t_name = t.get("task_name") or t.get("tast_name") or t.get("name") or "未命名任务"
                cleaned_tasks.append(
                    {
                        "task_name": t_name,
                        "description": t.get("description") or t.get("desc") or "",
                        "requirements": t.get("requirements") or t.get("reqs") or "",
                        "source_text_snippet": t.get("source_text_snippet") or "",
                    }
                )
            return cleaned_tasks
        except Exception as e:
            logger.error(f"AI 拆解任务失败: {e}")
            return [{"task_name": "解析失败", "description": str(e), "requirements": ""}]
