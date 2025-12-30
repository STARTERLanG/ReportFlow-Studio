from docx import Document


def inspect_table_structure():
    file_path = "docs/报告模板参考.docx"
    doc = Document(file_path)

    if not doc.tables:
        print("No tables found.")
        return

    table = doc.tables[0]
    print(
        f"Table found with {len(table.rows)} rows and {len(table.columns)} columns (grid)."
    )
    print("=" * 60)

    # Inspect first 50 rows to find patterns for "Section Headers"
    for i, row in enumerate(table.rows[:50]):
        # Get all cell texts
        cells = row.cells
        # Note: row.cells includes merged cells repeatedly.
        # We want to see unique cells to detect merging.
        unique_cells = []
        seen_ids = set()
        for cell in cells:
            if id(cell) not in seen_ids:
                unique_cells.append(cell)
                seen_ids.add(id(cell))

        texts = [c.text.strip() for c in unique_cells if c.text.strip()]
        full_text = " | ".join(texts)

        # Check heuristics for being a "Header Row"
        is_bold = False
        font_size = 0
        try:
            # Naive check of first run in first cell
            if unique_cells and unique_cells[0].paragraphs:
                run = unique_cells[0].paragraphs[0].runs[0]
                is_bold = run.bold
                if run.font.size:
                    font_size = run.font.size.pt
        except:
            pass

        # Print analysis
        prefix = f"Row {i:03d}"
        if len(texts) == 1:
            prefix += " [SINGLE COL]"  # Likely a merge-across header
        if is_bold:
            prefix += " [BOLD]"

        print(f"{prefix} | {full_text[:60]}...")


if __name__ == "__main__":
    inspect_table_structure()
