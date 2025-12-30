from docx import Document
from backend.app.services.template_service import TemplateService
import os


def create_sample_report(path):
    doc = Document()

    # 标题
    doc.add_heading("XX科技股份有限公司尽职调查报告", 0)

    # 第一章
    doc.add_heading("一、基本情况", 1)
    doc.add_paragraph(
        "XX科技股份有限公司（以下简称“公司”）成立于2010年，注册资本5000万元。公司主要从事智能硬件研发与销售。"
    )

    # 第二章
    doc.add_heading("二、财务分析", 1)
    doc.add_paragraph("近三年公司营收保持稳健增长，具体数据如下表所示：")

    # 表格
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "年度"
    table.cell(0, 1).text = "营业收入（万元）"
    table.cell(1, 0).text = "2023"
    table.cell(1, 1).text = "12,500"
    table.cell(2, 0).text = "2024"
    table.cell(2, 1).text = "15,800"

    doc.add_paragraph("从上表可以看出，公司年均复合增长率超过20%。")

    doc.save(path)


if __name__ == "__main__":
    test_file = "sample_report.docx"
    create_sample_report(test_file)

    service = TemplateService()
    result = service.parse_docx(test_file)

    print(f"\n样本解析结果 (Total Blocks: {result['total_blocks']}):")
    print("=" * 50)

    for block in result["blocks"]:
        type_tag = f"[{block.type.upper()}]"
        if block.type == "heading":
            type_tag += f" (H{block.level})"

        print(f"{type_tag:<15} | {block.original_text[:40]}...")
        if block.description:
            print(f"   Desc: {block.description}")
        print("-" * 50)

    if os.path.exists(test_file):
        os.remove(test_file)
